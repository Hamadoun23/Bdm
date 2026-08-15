"""
Exports avec graphiques Office natifs (objets modifiables, pas des images).

Portage de app/Services/GraphiquesDashboardExportService.php.

Côté Excel, openpyxl sait produire des graphiques natifs. Côté Word,
python-docx n'expose aucune API de graphique : la partie OOXML est donc écrite
à la main, comme le fait PhpWord, pour que l'utilisateur puisse toujours
modifier le graphique dans Word.
"""

import io

from docx import Document
from docx.shared import Pt
from django.http import HttpResponse
from openpyxl import Workbook
from openpyxl.chart import BarChart, DoughnutChart, PieChart, Reference

from .tableur import nom_onglet

#: Nombre de commerciaux et d'agences détaillés avant regroupement en « Autres ».
TOP_COMMERCIAUX = 5
TOP_AGENCES = 10


def libelle_volume(est_enrolement):
    return "Enrôlements" if est_enrolement else "Ventes"


def _regrouper(lignes, limite, libelle_reste, cle_label, denominateur=None):
    """
    Garde les `limite` premières lignes puis agrège la queue.

    Avec un `denominateur`, les valeurs sont converties en pourcentage.
    """
    avec_valeurs = sorted(
        [l for l in lignes if l["total_ventes"] > 0],
        key=lambda l: l["total_ventes"],
        reverse=True,
    )

    def valeur(total):
        return round(100 * total / denominateur, 2) if denominateur else int(total)

    labels = [l[cle_label] for l in avec_valeurs[:limite]]
    valeurs = [valeur(l["total_ventes"]) for l in avec_valeurs[:limite]]

    reste = avec_valeurs[limite:]
    if reste:
        labels.append(f"{libelle_reste} ({len(reste)})")
        valeurs.append(valeur(sum(l["total_ventes"] for l in reste)))

    return labels, valeurs


def donnees_graphiques_synthese(synthese):
    """Séries des trois graphiques de synthèse : types, commerciaux, agences."""
    par_type = synthese.get("par_type_carte") or []
    labels_type = [str(l["code"]) for l in par_type]
    valeurs_type = [int(l["total_ventes"]) for l in par_type]

    total = int((synthese.get("resume") or {}).get("total_ventes") or 0)
    labels_comm, valeurs_comm = _regrouper(
        synthese.get("commerciaux") or [],
        TOP_COMMERCIAUX,
        "Autres commerciaux",
        "user_name",
        denominateur=total or 1,
    )
    labels_ag, valeurs_ag = _regrouper(
        synthese.get("agences") or [], TOP_AGENCES, "Autres agences", "agence_nom"
    )

    return (
        {"labels": labels_type, "valeurs": valeurs_type},
        {"labels": labels_comm, "valeurs": valeurs_comm},
        {"labels": labels_ag, "valeurs": valeurs_ag},
    )


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------


def _remplir_deux_colonnes(feuille, entete_label, entete_valeur, labels, valeurs):
    feuille.cell(row=1, column=1, value=entete_label)
    feuille.cell(row=1, column=2, value=entete_valeur)
    for index, (label, valeur) in enumerate(zip(labels, valeurs), start=2):
        feuille.cell(row=index, column=1, value=label)
        feuille.cell(row=index, column=2, value=valeur)
    return len(labels)


def _references(feuille, nombre):
    """Plages de catégories et de valeurs du tableau à deux colonnes."""
    categories = Reference(feuille, min_col=1, min_row=2, max_row=nombre + 1)
    valeurs = Reference(feuille, min_col=2, min_row=1, max_row=nombre + 1)
    return categories, valeurs


def _ajouter_graphique(feuille, type_graphique, titre, ancre, nombre):
    if nombre == 0:
        return

    if type_graphique == "barres_horizontales":
        graphique = BarChart()
        graphique.type = "bar"
    elif type_graphique == "colonnes":
        graphique = BarChart()
        graphique.type = "col"
    elif type_graphique == "anneau":
        graphique = DoughnutChart()
    else:
        graphique = PieChart()

    categories, valeurs = _references(feuille, nombre)
    graphique.add_data(valeurs, titles_from_data=True)
    graphique.set_categories(categories)
    graphique.title = titre
    graphique.legend.position = "b"
    graphique.height = 12
    graphique.width = 22
    feuille.add_chart(graphique, ancre)


def classeur_synthese_campagne(nom_campagne, debut, fin, synthese, est_enrolement=False):
    """Classeur « graphiques » d'une synthèse de campagne."""
    par_type, commerciaux, agences = donnees_graphiques_synthese(synthese)
    libelle = libelle_volume(est_enrolement)

    classeur = Workbook()
    resume = classeur.active
    resume.title = nom_onglet("Resume")
    resume["A1"] = f"Synthèse — {nom_campagne}"
    resume["A2"] = f"{debut.strftime('%d/%m/%Y')} – {fin.strftime('%d/%m/%Y')}"
    resume["A3"] = (
        f"Total {libelle.lower()} : "
        f"{int((synthese.get('resume') or {}).get('total_ventes') or 0)}"
    )

    # Pas de type de carte sur une campagne d'enrôlement : l'onglet est omis.
    if not est_enrolement:
        feuille = classeur.create_sheet(nom_onglet("Types"))
        nombre = _remplir_deux_colonnes(
            feuille, "Type", "Ventes", par_type["labels"], par_type["valeurs"]
        )
        _ajouter_graphique(feuille, "anneau", "Mix types", "E2", nombre)

    feuille = classeur.create_sheet(nom_onglet("Commerciaux"))
    nombre = _remplir_deux_colonnes(
        feuille, "Commercial", "Part %", commerciaux["labels"], commerciaux["valeurs"]
    )
    _ajouter_graphique(feuille, "barres_horizontales", "Top commerciaux", "E2", nombre)

    feuille = classeur.create_sheet(nom_onglet("Agences"))
    nombre = _remplir_deux_colonnes(
        feuille, "Agence", libelle, agences["labels"], agences["valeurs"]
    )
    _ajouter_graphique(feuille, "secteurs", "Agences", "E2", nombre)

    classeur.active = 0
    return classeur


def classeur_performances(
    titre_periode, stats, top_commerciaux, ventes_par_agence, types_cartes, est_enrolement=False
):
    """Classeur « graphiques » de l'écran Performances."""
    libelle = libelle_volume(est_enrolement)

    classeur = Workbook()
    infos = classeur.active
    infos.title = nom_onglet("Infos")
    infos["A1"] = f"Performances — {titre_periode}"
    infos["A2"] = f"Total {libelle.lower()} : {int(stats.get('total_ventes') or 0)}"

    feuille = classeur.create_sheet(nom_onglet("Top commerciaux"))
    nombre = _remplir_deux_colonnes(
        feuille,
        "Commercial",
        libelle,
        [l["label"] for l in top_commerciaux],
        [int(l["ventes"]) for l in top_commerciaux],
    )
    _ajouter_graphique(feuille, "barres_horizontales", "Top commerciaux", "E2", nombre)

    feuille = classeur.create_sheet(nom_onglet("Agences"))
    nombre = _remplir_deux_colonnes(
        feuille,
        "Agence",
        libelle,
        [l["label"] for l in ventes_par_agence],
        [int(l["ventes"]) for l in ventes_par_agence],
    )
    _ajouter_graphique(feuille, "anneau", "Répartition agences", "E2", nombre)

    if not est_enrolement:
        par_type = stats.get("par_type") or {}
        feuille = classeur.create_sheet(nom_onglet("Types carte"))
        nombre = _remplir_deux_colonnes(
            feuille,
            "Type",
            "Ventes",
            [t.code for t in types_cartes],
            [int(par_type.get(str(t.id), par_type.get(t.id, 0)) or 0) for t in types_cartes],
        )
        _ajouter_graphique(feuille, "colonnes", "Ventes par type", "E2", nombre)

    classeur.active = 0
    return classeur


# ---------------------------------------------------------------------------
# Word — graphiques OOXML natifs
# ---------------------------------------------------------------------------

_ESPACES = (
    'xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" '
    'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
)

TYPE_CONTENU_GRAPHIQUE = (
    "application/vnd.openxmlformats-officedocument.drawingml.chart+xml"
)
RELATION_GRAPHIQUE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart"
)


def _echapper(texte):
    return (
        str(texte)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _xml_serie(labels, valeurs):
    """Catégories textuelles et valeurs numériques d'une série unique."""
    categories = "".join(
        f'<c:pt idx="{i}"><c:v>{_echapper(l)}</c:v></c:pt>'
        for i, l in enumerate(labels)
    )
    points = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(valeurs)
    )
    return (
        f'<c:cat><c:strRef><c:f>Feuil1!$A$2:$A${len(labels) + 1}</c:f>'
        f'<c:strCache><c:ptCount val="{len(labels)}"/>{categories}</c:strCache>'
        f"</c:strRef></c:cat>"
        f'<c:val><c:numRef><c:f>Feuil1!$B$2:$B${len(valeurs) + 1}</c:f>'
        f'<c:numCache><c:formatCode>General</c:formatCode>'
        f'<c:ptCount val="{len(valeurs)}"/>{points}</c:numCache>'
        f"</c:numRef></c:val>"
    )


def _xml_graphique(titre, type_graphique, labels, valeurs):
    """Document `chartSpace` complet pour un graphique à série unique."""
    serie = (
        f'<c:ser><c:idx val="0"/><c:order val="0"/>'
        f"<c:tx><c:strRef><c:f>Feuil1!$B$1</c:f><c:strCache>"
        f'<c:ptCount val="1"/><c:pt idx="0"><c:v>{_echapper(titre)}</c:v></c:pt>'
        f"</c:strCache></c:strRef></c:tx>"
        f"{_xml_serie(labels, valeurs)}</c:ser>"
    )

    if type_graphique in ("pie", "doughnut"):
        trou = '<c:holeSize val="50"/>' if type_graphique == "doughnut" else ""
        tracé = (
            f'<c:{"doughnutChart" if type_graphique == "doughnut" else "pieChart"}>'
            f'<c:varyColors val="1"/>{serie}{trou}'
            f'</c:{"doughnutChart" if type_graphique == "doughnut" else "pieChart"}>'
        )
        axes = ""
    else:
        direction = "bar" if type_graphique == "bar" else "col"
        tracé = (
            f'<c:barChart><c:barDir val="{direction}"/>'
            f'<c:grouping val="clustered"/><c:varyColors val="0"/>{serie}'
            f'<c:axId val="111111111"/><c:axId val="222222222"/></c:barChart>'
        )
        axes = (
            '<c:catAx><c:axId val="111111111"/><c:scaling><c:orientation val="minMax"/>'
            '</c:scaling><c:delete val="0"/><c:axPos val="l"/>'
            '<c:crossAx val="222222222"/></c:catAx>'
            '<c:valAx><c:axId val="222222222"/><c:scaling><c:orientation val="minMax"/>'
            '</c:scaling><c:delete val="0"/><c:axPos val="b"/>'
            '<c:crossAx val="111111111"/></c:valAx>'
        )

    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f"<c:chartSpace {_ESPACES}><c:chart>"
        f"<c:title><c:tx><c:rich><a:bodyPr/><a:lstStyle/><a:p><a:r>"
        f"<a:t>{_echapper(titre)}</a:t></a:r></a:p></c:rich></c:tx>"
        f'<c:overlay val="0"/></c:title><c:autoTitleDeleted val="0"/>'
        f"<c:plotArea><c:layout/>{tracé}{axes}</c:plotArea>"
        f'<c:legend><c:legendPos val="b"/><c:overlay val="0"/></c:legend>'
        f'<c:plotVisOnly val="1"/></c:chart></c:chartSpace>'
    ).encode("utf-8")


def _inserer_graphique(document, titre, type_graphique, labels, valeurs, largeur_cm=16, hauteur_cm=10):
    """Ajoute un graphique natif au document et retourne le paragraphe créé."""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    from docx.oxml import parse_xml

    partie_document = document.part
    numero = len(
        [r for r in partie_document.rels.values() if r.reltype == RELATION_GRAPHIQUE]
    ) + 1

    partie_graphique = Part(
        PackURI(f"/word/charts/chart{numero}.xml"),
        TYPE_CONTENU_GRAPHIQUE,
        _xml_graphique(titre, type_graphique, labels, valeurs),
        partie_document.package,
    )
    id_relation = partie_document.relate_to(partie_graphique, RELATION_GRAPHIQUE)

    largeur = int(largeur_cm * 360000)
    hauteur = int(hauteur_cm * 360000)
    paragraphe = document.add_paragraph()
    paragraphe.add_run()._r.append(
        parse_xml(
            '<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
            'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
            'xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f'<wp:inline><wp:extent cx="{largeur}" cy="{hauteur}"/>'
            f'<wp:docPr id="{numero}" name="Graphique {numero}"/><a:graphic>'
            '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/chart">'
            f'<c:chart r:id="{id_relation}"/></a:graphicData></a:graphic>'
            "</wp:inline></w:drawing>"
        )
    )
    return paragraphe


def document_graphiques(titre, sous_titres, graphiques):
    """
    Document Word : un titre, des lignes d'introduction, puis les graphiques.

    `graphiques` est une liste de `(titre, type, labels, valeurs)` où `type`
    vaut « bar », « column », « pie » ou « doughnut ».
    """
    document = Document()

    entete = document.add_paragraph()
    passage = entete.add_run(titre)
    passage.bold = True
    passage.font.size = Pt(16)

    for ligne in sous_titres:
        document.add_paragraph(ligne)
    document.add_paragraph()

    for titre_graphique, type_graphique, labels, valeurs in graphiques:
        paragraphe = document.add_paragraph()
        passage = paragraphe.add_run(titre_graphique)
        passage.bold = True
        passage.font.size = Pt(12)
        _inserer_graphique(document, titre_graphique, type_graphique, labels, valeurs)
        document.add_paragraph()

    return document


def reponse_docx(document, nom_fichier):
    if not nom_fichier.lower().endswith(".docx"):
        nom_fichier += ".docx"

    tampon = io.BytesIO()
    document.save(tampon)
    tampon.seek(0)

    reponse = HttpResponse(
        tampon.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    reponse["Content-Disposition"] = f'attachment; filename="{nom_fichier}"'
    reponse["Cache-Control"] = "max-age=0"
    return reponse
