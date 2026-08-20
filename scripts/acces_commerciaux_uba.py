#!/usr/bin/env python
"""
Pose les mots de passe des commerciaux UBA et produit leur fiche d'accès Word.

Reprend la convention retenue pour la campagne BDM d'août 2026 :

    initiale du prénom + deux derniers chiffres du téléphone
    + initiale du nom + « @uba »

    Fatou MAIGA, 76208554  →  F54M@uba

Chaque mot de passe est donc propre à son porteur et ne fonctionne qu'avec son
numéro : une ligne recopiée par erreur ne donne accès à rien. Le suffixe change
de client (« @bdm » pour la BDM) pour qu'un commercial des deux campagnes ne
confonde pas ses deux accès.

Le document reprend la mise en forme du fichier BDM, sans la colonne « Agence » :
UBA n'a pas de réseau d'agences.

Sur toute base autre que `bdm_dev`, il exige `--production` : le script écrit
des mots de passe, on ne le déclenche pas par inadvertance.

Usage :
    backend/.venv/Scripts/python.exe scripts/acces_commerciaux_uba.py
    python scripts/acces_commerciaux_uba.py --production
"""

import argparse
import os
import sys
from datetime import date

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings")

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

import django  # noqa: E402

django.setup()

from django.conf import settings  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.table import WD_ALIGN_VERTICAL  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt, RGBColor  # noqa: E402

from core.auth_backend import hacher_mot_de_passe  # noqa: E402
from core.models import Partenaire, Role, User  # noqa: E402

BASE_DEVELOPPEMENT = "bdm_dev"

#: Suffixe du mot de passe — il désigne le client, pas la banque émettrice.
SUFFIXE = "@uba"

DOSSIER_SORTIE = os.path.join(os.path.dirname(__file__), "..", "docs", "UBA")
NOM_FICHIER = "Acces Campagne GDA-UBA Aout 2026.docx"

# Couleurs reprises du document BDM.
BORDEAUX = RGBColor(0x38, 0x14, 0x19)
ORANGE = RGBColor(0xE8, 0x5D, 0x2A)
GRIS = RGBColor(0x59, 0x59, 0x59)
GRIS_CLAIR = RGBColor(0x8C, 0x8C, 0x8C)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Mots de passe
# ---------------------------------------------------------------------------


def mot_de_passe(user) -> str:
    """
    `F54M@uba` — initiale du prénom, fin du numéro, initiale du nom.

    Le numéro est nettoyé de tout séparateur : certaines fiches portent des
    espaces ou des tirets, et le mot de passe doit rester déductible de ce qui
    est imprimé dans la colonne « Identifiant ».
    """
    prenom = (user.prenom or "").strip()
    nom = (user.name or "").strip()
    chiffres = "".join(c for c in (user.telephone or "") if c.isdigit())

    initiale_prenom = prenom[0].upper() if prenom else "X"
    initiale_nom = nom[0].upper() if nom else "X"
    return f"{initiale_prenom}{chiffres[-2:]}{initiale_nom}{SUFFIXE}"


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------


def _fond(cellule, couleur_hex):
    """Applique une couleur de fond à une cellule (pas d'API python-docx pour ça)."""
    ombrage = OxmlElement("w:shd")
    ombrage.set(qn("w:val"), "clear")
    ombrage.set(qn("w:color"), "auto")
    ombrage.set(qn("w:fill"), couleur_hex)
    cellule._tc.get_or_add_tcPr().append(ombrage)


def _bordures(tableau, couleur_hex="D9D9D9", epaisseur=6):
    bordures = OxmlElement("w:tblBorders")
    for cote in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{cote}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(epaisseur))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), couleur_hex)
        bordures.append(element)
    tableau._tbl.tblPr.append(bordures)


def _paragraphe(document, texte, taille=None, gras=False, couleur=None, apres=Pt(2)):
    paragraphe = document.add_paragraph()
    paragraphe.paragraph_format.space_after = apres
    serie = paragraphe.add_run(texte)
    if taille:
        serie.font.size = taille
    serie.font.bold = gras
    if couleur:
        serie.font.color.rgb = couleur
    return paragraphe


def _cellule(cellule, texte, gras=False, couleur=None, centre=False):
    cellule.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraphe = cellule.paragraphs[0]
    paragraphe.paragraph_format.space_after = Pt(0)
    if centre:
        paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    serie = paragraphe.add_run(texte)
    serie.font.bold = gras
    if couleur:
        serie.font.color.rgb = couleur


def construire_document(commerciaux, site, campagne_nom, periode):
    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.6)

    _paragraphe(
        document,
        "Accès à l’application — Campagne GDA/UBA Août 2026",
        taille=Pt(18), gras=True, couleur=BORDEAUX, apres=Pt(3),
    )
    _paragraphe(
        document,
        f"Vente des cartes VISA prépayées GDA · {periode}",
        couleur=GRIS, apres=Pt(12),
    )

    _paragraphe(
        document, "Comment se connecter",
        taille=Pt(12), gras=True, couleur=ORANGE, apres=Pt(5),
    )
    for etape in (
        f"Ouvrir le site : {site}",
        "Saisir son numéro de téléphone dans le champ « Identifiant » "
        "(chiffres uniquement, sans espace).",
        "Saisir le mot de passe indiqué sur SA PROPRE ligne du tableau ci-dessous.",
    ):
        _paragraphe(document, etape)
    _paragraphe(
        document,
        "Le mot de passe respecte les majuscules et minuscules : la 1re lettre "
        "et la dernière lettre sont en MAJUSCULE, la fin « @uba » est en "
        "minuscules.",
        apres=Pt(12),
    )

    _paragraphe(
        document, "Identifiants par commercial",
        taille=Pt(12), gras=True, couleur=ORANGE, apres=Pt(5),
    )
    _paragraphe(
        document,
        "Chaque commercial a un mot de passe qui lui est propre : il ne "
        "fonctionne qu’avec son propre numéro. En cas d’échec de connexion, "
        "vérifier que la ligne utilisée est bien la sienne, puis contacter "
        "l’administrateur.",
        taille=Pt(9), couleur=GRIS, apres=Pt(4),
    )

    # Trois colonnes seulement : UBA n'a pas d'agences, ses commerciaux
    # dépendent directement du client.
    tableau = document.add_table(rows=1, cols=3)
    _bordures(tableau)
    for cellule, intitule in zip(
        tableau.rows[0].cells, ("Nom", "Identifiant (téléphone)", "Mot de passe")
    ):
        _fond(cellule, "381419")
        _cellule(cellule, intitule, gras=True, couleur=BLANC)

    for user, secret in commerciaux:
        ligne = tableau.add_row().cells
        _cellule(ligne[0], user.nom_complet)
        _cellule(ligne[1], user.telephone or "—")
        _cellule(ligne[2], secret, gras=True, couleur=BORDEAUX)

    for cellule_largeur, largeur in zip(tableau.columns, (Cm(8.0), Cm(4.5), Cm(4.5))):
        for cellule in cellule_largeur.cells:
            cellule.width = largeur

    _paragraphe(document, "", apres=Pt(4))
    _paragraphe(
        document,
        f"Document généré le {date.today():%d/%m/%Y} — {len(commerciaux)} "
        f"commerciaux, campagne « {campagne_nom} ».",
        taille=Pt(8), couleur=GRIS_CLAIR, apres=Pt(0),
    )
    return document


# ---------------------------------------------------------------------------


def main():
    analyseur = argparse.ArgumentParser(description=__doc__)
    analyseur.add_argument(
        "--site",
        default="bdm.gdamali.net",
        help="adresse imprimée sur la fiche (défaut : le site actuel)",
    )
    analyseur.add_argument(
        "--sans-ecriture",
        action="store_true",
        help="génère le document sans toucher aux mots de passe en base",
    )
    analyseur.add_argument(
        "--production",
        action="store_true",
        help="autorise l'exécution sur une base autre que bdm_dev",
    )
    options = analyseur.parse_args()

    base = settings.DATABASES["default"]["NAME"]
    if (
        base != BASE_DEVELOPPEMENT
        and not options.sans_ecriture
        and not options.production
    ):
        raise SystemExit(
            f"REFUS : base « {base} » (hors développement). Ce script écrit des "
            "mots de passe : relancez avec --production."
        )

    partenaire = Partenaire.objects.filter(code="uba").first()
    if partenaire is None:
        raise SystemExit("Le partenaire UBA est absent : lancez d'abord les migrations.")

    users = list(
        User.objects.filter(partenaire_id=partenaire.id, role=Role.COMMERCIAL)
        .order_by("name", "prenom")
    )
    if not users:
        raise SystemExit(
            "Aucun commercial UBA : lancez `scripts/preparer_campagne_uba.py`."
        )

    commerciaux = [(user, mot_de_passe(user)) for user in users]

    doublons = {s for _, s in commerciaux if [x for _, x in commerciaux].count(s) > 1}
    if doublons:
        raise SystemExit(
            f"Collision de mots de passe : {sorted(doublons)}. "
            "Deux commerciaux partageraient le même accès — corrigez la convention."
        )

    if not options.sans_ecriture:
        for user, secret in commerciaux:
            user.password = hacher_mot_de_passe(secret)
            user.save(update_fields=["password"])

    campagne = partenaire.campagnes.order_by("-date_debut").first()
    periode = (
        f"du {campagne.date_debut:%d/%m/%Y} au {campagne.date_fin:%d/%m/%Y}"
        if campagne
        else "période à définir"
    )
    document = construire_document(
        commerciaux, options.site, campagne.nom if campagne else "—", periode
    )

    os.makedirs(DOSSIER_SORTIE, exist_ok=True)
    chemin = os.path.abspath(os.path.join(DOSSIER_SORTIE, NOM_FICHIER))
    document.save(chemin)

    print(f"Base       : {base}")
    print(f"Client     : {partenaire.nom} — {partenaire.nom_complet}")
    print(f"Campagne   : {campagne.nom if campagne else '—'} ({periode})")
    print(f"Site       : {options.site}")
    print(
        "Écriture   : "
        + ("aucune (--sans-ecriture)" if options.sans_ecriture else "mots de passe posés")
    )
    print(f"Document   : {chemin}\n")

    print(f"{'Nom':<28} {'Téléphone':<12} Mot de passe")
    print("-" * 54)
    for user, secret in commerciaux:
        print(f"{user.nom_complet:<28} {user.telephone or '—':<12} {secret}")


if __name__ == "__main__":
    main()
